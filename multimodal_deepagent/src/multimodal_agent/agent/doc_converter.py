"""Convert uploaded document bytes to markdown text for the agent's virtual filesystem."""

from __future__ import annotations

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

DOCUMENT_MIMES: frozenset[str] = frozenset({
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "text/plain",
    "text/markdown",
    "text/csv",
})

DOCUMENT_EXTENSIONS: frozenset[str] = frozenset({
    ".pdf", ".docx", ".doc", ".txt", ".md", ".csv",
})


def is_document(mime: str, filename: str = "") -> bool:
    """Return True if the MIME type or file extension indicates a document."""
    if mime in DOCUMENT_MIMES:
        return True
    ext = Path(filename).suffix.lower()
    return ext in DOCUMENT_EXTENSIONS


def virtual_filename(original_name: str) -> str:
    """Return the /uploads/ stem used to store the converted file.

    Plain-text formats keep their original extension; everything else gets .md.
    """
    p = Path(original_name)
    if p.suffix.lower() in {".txt", ".md", ".csv"}:
        return p.name
    return p.stem + ".md"


def convert_to_markdown(raw_bytes: bytes, mime: str, filename: str) -> str:
    """Convert *raw_bytes* to a UTF-8 markdown string.

    Dispatches by MIME type, falling back to file extension when the MIME is
    generic (e.g. ``application/octet-stream``).
    """
    ext = Path(filename).suffix.lower()

    if mime == "application/pdf" or ext == ".pdf":
        return _pdf_to_text(raw_bytes, filename)

    if mime in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or ext in (".docx", ".doc"):
        return _docx_to_markdown(raw_bytes, filename)

    # Plain text / markdown / CSV — decode as UTF-8
    return raw_bytes.decode("utf-8", errors="replace")


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _pdf_to_text(raw_bytes: bytes, filename: str) -> str:
    try:
        from pdfminer.high_level import extract_text  # type: ignore[import]
        text = extract_text(io.BytesIO(raw_bytes))
        if text and text.strip():
            return text.strip()
        return f"[{filename}: no extractable text — may be a scanned/image PDF]"
    except ImportError:
        logger.warning("pdfminer.six not installed; cannot extract PDF text")
        return f"[{filename}: pdfminer.six not installed]"
    except Exception as exc:
        logger.warning("PDF extraction failed for %s: %s", filename, exc)
        return f"[{filename}: extraction failed — {exc}]"


# ---------------------------------------------------------------------------
# DOCX / DOC
# ---------------------------------------------------------------------------

def _docx_to_markdown(raw_bytes: bytes, filename: str) -> str:
    try:
        from docx import Document  # type: ignore[import]
        doc = Document(io.BytesIO(raw_bytes))
        lines: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            style = para.style.name if para.style else ""
            if not text:
                lines.append("")
                continue
            if style.startswith("Heading 1"):
                lines.append(f"# {text}")
            elif style.startswith("Heading 2"):
                lines.append(f"## {text}")
            elif style.startswith("Heading 3"):
                lines.append(f"### {text}")
            elif style.startswith("Heading"):
                lines.append(f"#### {text}")
            else:
                lines.append(text)

        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            col_count = max(len(r) for r in rows)
            lines.append("")
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("| " + " | ".join(["---"] * col_count) + " |")
            for row in rows[1:]:
                padded = row + [""] * (col_count - len(row))
                lines.append("| " + " | ".join(padded) + " |")
            lines.append("")

        return "\n".join(lines).strip()

    except ImportError:
        logger.warning("python-docx not installed; cannot extract DOCX text")
        return f"[{filename}: python-docx not installed]"
    except Exception as exc:
        logger.warning("DOCX extraction failed for %s: %s", filename, exc)
        return f"[{filename}: extraction failed — {exc}]"
